from docx import Document
from docx.shared import Inches
from beam import *


###
class GeneralCalculatorForTwoHorizontal:

    def __init__(self):
        print("calculatorOneHorizontal Initialized")
        self.calculatedValues = CalculatedValues()
        self.verticalHatil = None
        self.horizontalHatilBottom = None
        self.horizontalHatilTop = None
        self.concrete = None
        self.steel = None
        self.wall = None
        self.plaster = None
        self.earthquake = None
        self.reinforcedConcreteDensity = None
        self.concreteCover = None
        self.heightParameter = None
        self.report = ""

    def calculateTwoHorizontal(self, verticalHatil, horizontalHatilBottom, horizontalHatilTop, concrete, steel, wall,
    plaster, earthquake, reinforcedConcreteDensity, concreteCover, heightParameter):
        self.verticalHatil = verticalHatil
        self.horizontalHatilBottom = horizontalHatilBottom
        self.horizontalHatilTop = horizontalHatilTop
        self.concrete = concrete
        self.steel = steel
        self.wall = wall
        self.plaster = plaster
        self.earthquake = earthquake
        self.reinforcedConcreteDensity = reinforcedConcreteDensity
        self.concreteCover = concreteCover
        self.heightParameter = heightParameter

document = Document()

document.add_heading('HatilApp')

pic = document.add_picture('word/logo.png', width=Inches(2.5))

#I can write the parameter(integer) in the text.
x=1

p = document.add_paragraph()
p.add_run('DOCUMENT NO      :').bold = True
p.add_run(" HatilApp " + str(x))
p = document.add_paragraph()
p.add_run('DOCUMENT TITLE   :').bold = True
p = document.add_paragraph()
p.add_run('REVISION NO      :').bold = True
p = document.add_paragraph()
p.add_run('DATE             :').bold = True

document.add_heading('HatilApp Basement Floor Hatil Report', level=0)

document.add_heading('Weight per unit area of wall', level=1)

p = document.add_paragraph()
p.add_run("Wall's density = ").bold = True
p.add_run(str(Wall.density))





document.add_page_break()

document.save('word/wordDoc.docx')






verticalHatil = VerticalHatil(20, 4, 7)
horizontalHatilBottom = HorizontalHatil(20, 2.3)
horizontalHatilTop = HorizontalHatil(20, 4.7)
concrete = Concrete("C25")
steel = ReinforcementSteel("S420", "ø10", "ø8", 20)
wall = Wall(25, 0.5, 8)
plaster = Plaster(2, 1.8)
earthquake = Earthquake(0.4, 1)
reinforcedConcreteDensity = ReinforcedConcreteDensity()
concreteCover = ConcreteCover(3)
heightParameter = HeightParameter(-5, 10)


calculatorTwoHorizontal = GeneralCalculatorForTwoHorizontal()
calculatorTwoHorizontal.calculateTwoHorizontal(verticalHatil, horizontalHatilBottom, horizontalHatilTop, concrete, steel, wall, plaster,
earthquake, reinforcedConcreteDensity, concreteCover, heightParameter)